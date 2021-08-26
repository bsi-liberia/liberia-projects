import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def looks_like_num(value):
    try:
        float(re.sub(",", "", value))
        return True
    except ValueError:
        return False


def create_document():
    document = Document()
    section = document.sections[-1]
    section.page_height = Mm(297)  # for A4 paper
    section.page_width = Mm(210)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    return document


def create_section(document, orientation="landscape"):
    if orientation == "portrait":
        document.add_section()
        section = document.sections[-1]
        section.page_width = Mm(297)  # for A4 paper
        section.page_height = Mm(210)
    else:
        document.add_section()
        section = document.sections[-1]
        section.page_width = Mm(297)  # for A4 paper
        section.page_height = Mm(210)
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        section.left_margin = Mm(12.7)
        section.right_margin = Mm(12.7)
        return document


def add_heading(i, value, hdr_cells):
    hdr_cells[i].text = ''
    hdr_cells[i].paragraphs[0].add_run(value).bold = True
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return hdr_cells


def add_cell_value(i, value, cells, html_row, to_right=False):
    cells[i].text = value
    if to_right:
        cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return cells


def html_table_to_doc(html_table, document):
    return HTMLTableDoc(html_table, document).document


class RowsTableDoc(object):
    def __init__(self, rows, document, headers=None):
        self.document = document
        self.rows = rows

        if (type(rows[0]) in (list, tuple)):
            row_cols = len(rows[0])
        elif type(rows[0]) == dict:
            row_cols = len(rows[0].keys())
        else:
            print("it's a different type", type(rows[0]))

        self.table = self.document.add_table(rows=0, cols=row_cols)
        self.table.autofit = True
        self.table.style = 'Grid Table 4 - Accent 11'
        self.rowspans = {}

        if headers:
            if len(headers) != row_cols:
                raise Exception(
                    "Headers, if specified, must be the same length as the number of keys in the rows.")
            hdr_cells = self.table.add_row().cells
            for i, header in enumerate(headers):
                hdr_cells = add_heading(i, header, hdr_cells)
        self.add_rows()

    def add_rows(self):
        for row in self.rows:
            if type(row) in (list, tuple):
                _row = row
            elif type(row) == dict:
                _row = row.values()
            row_cells = self.table.add_row().cells
            for i, cell in enumerate(_row):
                row_cells = add_cell_value(i, cell, row_cells, _row)

    def document(self):
        return self.document


def rows_to_table(rows, document, headers=None):
    return RowsTableDoc(rows, document, headers).document
